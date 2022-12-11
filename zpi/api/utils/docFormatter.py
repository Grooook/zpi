import os

from docx import Document
from docxtpl import DocxTemplate
import re

from api.models import ApplicationProperty, Property, UserApplication, UserApplicationProperty

from api.utils.utils import get_property_dict
from api.serializers import UserApplicationSerializer



class DocFormatter:
    def __init__(self, application):
        self.regex = re.compile(r"(?<={{)(.*?)(?=}})")
        self.application = application
        self.file_path = application.file.path
        self.document = Document(self.file_path)
        self.params = self.get_params()

    def save_to_db(self):
        properties = get_property_dict()
        for param in self.params:
            property = Property.objects.get(pk=properties[param])
            ApplicationProperty.objects.get_or_create(application=self.application, property=property)


    def get_params(self):
        params = []
        for p in self.document.paragraphs:
            if self.regex.findall(p.text):
                params += (self.regex.findall(p.text))

        for t in self.document.tables:
            for t_row in t.rows:
                for r_cell in t_row.cells:
                    for c_paragraph in r_cell.paragraphs:
                        if self.regex.findall(c_paragraph.text):
                            params += (self.regex.findall(c_paragraph.text))

        return params

    def save_new_document(self, user_application, context):
        if 'csrfmiddlewaretoken' in context:
            del context['csrfmiddlewaretoken']
        context = {key: context[key][0] for key in context.keys()}
        document = DocxTemplate(self.file_path)
        document.render(context=context)
        new_path = "documents/processed/" + self.application.file.name.split('/')[-1]
        document.save("media/" + new_path)
        user_application = UserApplication.objects.get(pk=user_application)
        user_application.file = new_path
        user_application.save()


    def update_document(self, user_application):
        user_application_properties = UserApplicationProperty.objects.filter(user_application__pk=user_application).values('property__name', 'value')
        context = {property['property__name']: property['value'] for property in user_application_properties}
        document = DocxTemplate(self.file_path)
        document.render(context=context)
        new_path = "documents/processed/" + self.application.file.name.split('/')[-1]
        if os.path.exists("media/" + new_path):
            os.remove("media/" + new_path)
        document.save("media/" + new_path)
        user_application = UserApplication.objects.get(pk=user_application)
        user_application.file = new_path
        user_application.save()
